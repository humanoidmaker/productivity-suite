"""Tests for DOCX exporter."""
import io
from docx import Document
from app.export.docx_exporter import export_document


def _open(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def test_export_heading():
    html = "<h1>Title</h1>"
    doc = _open(export_document(html, "Test"))
    assert any(p.text == "Title" for p in doc.paragraphs)

def test_export_paragraph():
    html = "<p>Hello world</p>"
    doc = _open(export_document(html, "Test"))
    assert any("Hello world" in p.text for p in doc.paragraphs)

def test_export_bold():
    html = "<p><strong>Bold text</strong></p>"
    doc = _open(export_document(html, "Test"))
    for p in doc.paragraphs:
        for run in p.runs:
            if "Bold text" in run.text:
                assert run.bold

def test_export_bullet_list():
    html = "<ul><li>Item 1</li><li>Item 2</li></ul>"
    data = export_document(html, "Test")
    doc = _open(data)
    texts = [p.text for p in doc.paragraphs]
    assert "Item 1" in texts
    assert "Item 2" in texts

def test_export_empty():
    data = export_document("", "Empty")
    doc = _open(data)
    assert len(doc.paragraphs) > 0  # Should have at least the "(Empty document)" text

def test_export_valid_docx():
    html = "<h1>Test</h1><p>Content</p>"
    data = export_document(html, "Test")
    assert len(data) > 100
    # Verify it's a valid DOCX (ZIP with correct structure)
    import zipfile
    zf = zipfile.ZipFile(io.BytesIO(data))
    assert "word/document.xml" in zf.namelist()

def test_export_hr():
    html = "<hr>"
    doc = _open(export_document(html, "Test"))
    assert any("─" in p.text for p in doc.paragraphs)

def test_export_multiple_paragraphs():
    html = "<p>One</p><p>Two</p><p>Three</p>"
    doc = _open(export_document(html, "Test"))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert len(texts) >= 3

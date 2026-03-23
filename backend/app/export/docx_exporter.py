from __future__ import annotations
"""Export HTML/TipTap content to DOCX using python-docx."""

import io
import re
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class _HTMLToDocx(HTMLParser):
    """Parse HTML and build DOCX paragraphs."""

    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.bold = False
        self.italic = False
        self.underline = False
        self.in_list = False
        self.list_type = "bullet"
        self.in_code = False
        self.heading_level = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attr_dict = dict(attrs)
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.heading_level = int(tag[1])
            self.current_paragraph = self.doc.add_heading("", level=self.heading_level)
        elif tag == "p":
            self.current_paragraph = self.doc.add_paragraph()
            align = attr_dict.get("style", "")
            if "text-align: center" in align:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "text-align: right" in align:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif tag in ("ul", "ol"):
            self.in_list = True
            self.list_type = "bullet" if tag == "ul" else "number"
        elif tag == "li":
            style = "List Bullet" if self.list_type == "bullet" else "List Number"
            self.current_paragraph = self.doc.add_paragraph(style=style)
        elif tag == "strong" or tag == "b":
            self.bold = True
        elif tag == "em" or tag == "i":
            self.italic = True
        elif tag == "u":
            self.underline = True
        elif tag == "pre" or tag == "code":
            self.in_code = True
            if tag == "pre":
                self.current_paragraph = self.doc.add_paragraph()
        elif tag == "br":
            if self.current_paragraph:
                self.current_paragraph.add_run("\n")
        elif tag == "hr":
            self.doc.add_paragraph("─" * 50)
        elif tag == "blockquote":
            self.current_paragraph = self.doc.add_paragraph()
            pf = self.current_paragraph.paragraph_format
            pf.left_indent = Inches(0.5)
        elif tag == "table":
            pass  # Tables handled separately
        elif tag == "img":
            pass  # Images handled separately

    def handle_endtag(self, tag: str) -> None:
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.heading_level = 0
            self.current_paragraph = None
        elif tag == "p":
            self.current_paragraph = None
        elif tag in ("ul", "ol"):
            self.in_list = False
        elif tag == "li":
            self.current_paragraph = None
        elif tag in ("strong", "b"):
            self.bold = False
        elif tag in ("em", "i"):
            self.italic = False
        elif tag == "u":
            self.underline = False
        elif tag in ("pre", "code"):
            self.in_code = False

    def handle_data(self, data: str) -> None:
        text = data
        if not text.strip() and not self.in_code:
            return
        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph()
        run = self.current_paragraph.add_run(text)
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline
        if self.in_code:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        if self.heading_level == 0:
            run.font.size = Pt(11)


def export_document(html: str, title: str) -> bytes:
    """Convert HTML content to DOCX bytes."""
    doc = Document()
    doc.core_properties.title = title

    if not html or not html.strip():
        doc.add_paragraph("(Empty document)")
    else:
        parser = _HTMLToDocx(doc)
        parser.feed(html)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
